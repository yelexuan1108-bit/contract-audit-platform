# -*- coding: utf-8 -*-
"""Generate detailed EN + ZH user guide .docx for the contract-audit platform."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = "docs"
EN_PATH = f"{OUT_DIR}/USER_GUIDE_EN.docx"
ZH_PATH = f"{OUT_DIR}/USER_GUIDE_ZH.docx"

ACCENT = RGBColor(0x4F, 0x46, 0xE5)   # indigo
DARK = RGBColor(0x1F, 0x29, 0x37)
GREY = RGBColor(0x6B, 0x72, 0x80)


def set_base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.font.color.rgb = DARK
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.15


def set_cn_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.font.italic = italic


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 0:
        set_cn_font(run, size=22, bold=True, color=ACCENT)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
    elif level == 1:
        set_cn_font(run, size=16, bold=True, color=ACCENT)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "4F46E5")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        set_cn_font(run, size=13, bold=True, color=DARK)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    else:
        set_cn_font(run, size=11.5, bold=True, color=GREY)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    return p


def add_para(doc, text, bold=False, color=None, size=None, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cn_font(run, size=size, bold=bold, color=color, italic=italic)
    return p


def add_bullets(doc, items, level=0):
    for it in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        if isinstance(it, tuple):
            head, rest = it
            run = p.add_run(head)
            set_cn_font(run, bold=True)
            run2 = p.add_run(rest)
            set_cn_font(run2)
        else:
            run = p.add_run(it)
            set_cn_font(run)


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        set_cn_font(run, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "4F46E5")
        hdr[i]._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            set_cn_font(run)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    set_cn_font(run)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F3F4F6")
    p._p.get_or_add_pPr().append(shd)
    p.paragraph_format.left_indent = Inches(0.15)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cn_font(run, italic=True, color=RGBColor(0xB4, 0x5F, 0x06))
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "FEF3C7")
    p._p.get_or_add_pPr().append(shd)
    p.paragraph_format.left_indent = Inches(0.15)
    return p


# =====================================================================
# ENGLISH
# =====================================================================
def build_en():
    doc = Document()
    set_base_style(doc)

    add_heading(doc, "Contract Note Audit Platform — User Guide", 0)
    add_para(doc, "An AI-assisted audit platform for Hong Kong investment contract notes. "
                  "It compares each contract note against official templates and reports missing or inconsistent clauses.",
             color=GREY)

    # 1 Overview --------------------------------------------------------
    add_heading(doc, "1. Overview", 1)
    add_para(doc, "The platform audits four product types. Each product has its own page and its own "
                  "auto-detection and clause-checking logic.")
    add_table(doc,
              ["Product", "Page", "Description"],
              [
                  ["Stock (股票)", "Stock Contract Note", "Equity trades; auto-detects HK/US market and one of 6 subtypes."],
                  ["Fund (基金)", "Fund Contract Note", "Fund transactions; auto-detects language and one of 6 transaction subtypes."],
                  ["VA (虚拟资产)", "VA Virtual Asset", "Crypto contract notes; auto-detects Buy/Sell and language; runs a 50-item VA-specific checker."],
                  ["Statement (月结单)", "Investment Monthly Statement", "Monthly statements; 3 subtypes: Monthly / BB Invest / Southbound."],
              ],
              widths=[1.5, 1.9, 3.1])
    add_para(doc, "")
    add_para(doc, "Every audit performs the following checks:", bold=True)
    add_bullets(doc, [
        ("Clause audit (条款核查) — ", "a local rule engine compares PDF text against stored clause templates and required fields."),
        ("AI audit (AI 审核) — ", "an optional LLM review (Stock and Fund only) that flags risk findings, fee waivers and a summary."),
        ("VA special check (VA 专项检查) — ", "a dedicated 50-item checker for Virtual Asset contract notes."),
    ])

    # 2 Setup -----------------------------------------------------------
    add_heading(doc, "2. Setup", 1)
    add_heading(doc, "2.1 Requirements", 2)
    add_bullets(doc, [
        "Python 3.10 or later.",
        "Dependencies listed in requirements.txt (FastAPI, uvicorn, pdfplumber, PyPDF2, pydantic, beautifulsoup4, openpyxl, python-docx, httpx).",
    ])
    add_heading(doc, "2.2 Install & run", 2)
    add_code(doc, "pip install -r requirements.txt\ncp .env.example .env    # then edit .env\npython app.py")
    add_para(doc, "Open http://localhost:8000 in a browser.")
    add_heading(doc, "2.3 Configuration (.env)", 2)
    add_table(doc,
              ["Variable", "Purpose", "Required"],
              [
                  ["OPENROUTER_API_KEY", "API key for the AI audit (OpenRouter/Claude). If unset, only the local rule engine runs.", "No"],
                  ["PORT", "Server port (default 8000).", "No"],
              ],
              widths=[1.8, 3.8, 0.9])
    add_para(doc, "")
    add_para(doc, "A badge in the top-right corner shows whether AI is enabled. Without an API key, "
                  "clause auditing and VA checks still run fully offline.")

    # 3 Auditing a product ----------------------------------------------
    add_heading(doc, "3. Auditing a Product", 1)
    add_para(doc, "All product pages (Stock / Fund / VA / Statement) follow the same workflow.")
    add_heading(doc, "Step 1 — Upload files", 2)
    add_bullets(doc, [
        "Drag a folder or PDF files into the drop zone, or click 选择 PDF 文件 / 选择文件夹.",
        "Multiple files are supported (up to 200 per batch).",
        "The file list shows each selected file and the total size.",
    ])
    add_heading(doc, "Step 2 — (Statement only) choose a subtype", 2)
    add_para(doc, "On the Statement page select 子类型: 投资月结单 (monthly) / BB Invest / Southbound 南向通.")
    add_heading(doc, "Step 3 — (Stock / Fund only) optionally skip AI", 2)
    add_para(doc, "Tick 跳过 AI 审核 to run only the local rule engine. This is much faster for large batches.")
    add_heading(doc, "Step 4 — Run", 2)
    add_para(doc, "Click 开始审核. A progress bar shows live upload/processing progress. Files are processed "
                  "concurrently (up to 10 workers).")
    add_heading(doc, "Step 5 — Review results", 2)
    add_para(doc, "Each file is tagged with its detected product, subtype and overall result:")
    add_bullets(doc, [
        ("条款通过 (clauses passed) — ", "all clause checks passed."),
        ("条款问题 (clause issues) — ", "lists the failing clauses."),
    ])
    add_para(doc, "Expand a result to see:", bold=True)
    add_bullets(doc, [
        ("English / Chinese clauses (英文条款 / 中文条款) — ", "per-clause pass/fail with matched keywords."),
        ("Personal information (个人资料) — ", "fixed field-label check."),
        ("AI audit (AI 审核) — ", "findings, critical/warning/info counts and summary (Stock/Fund only)."),
        ("VA special check (VA 专项检查) — ", "the 50-item VA checklist (VA only)."),
    ])

    add_heading(doc, "3.1 Auto-detection of file type", 2)
    add_para(doc, "The auditor identifies the document by keywords, so files placed on the wrong page are "
                  "handled according to content (and the correct page is still recommended).")
    add_bullets(doc, [
        ("Stock subtypes — ", "order (US buy/sell), orderHk (HK buy/sell), companyAction / companyActionHk "
                              "(Corporate Action Advice), business (Money/Stock Transfer / Deposit), "
                              "businessHk (IPO allotment)."),
        ("Fund subtypes — ", "apply (Subscription), redeem (Redemption), deposit, takeout (Withdrawal), "
                              "divid (Cash Dividend), divid_share (Unit Dividend)."),
        ("VA — ", "Buy vs Sell and language (Traditional / Simplified Chinese)."),
        ("Statement — ", "Southbound detected by keyword; the rest default to monthly."),
    ])

    add_heading(doc, "3.2 Date validation", 2)
    add_para(doc, "The engine validates that Contract/Dealing Date plus N Hong Kong business days equals the "
                  "Issue Date, using the HK public-holiday calendar (2024–2027, weekends + public holidays).")
    add_bullets(doc, [
        ("Stock / VA — ", "allow T, T+1 or T+2 business days."),
        ("Fund / Statement — ", "strict T+2 business days."),
    ])

    # 4 Clause template management -------------------------------------
    add_heading(doc, "4. Clause Template Management (条款模板)", 1)
    add_para(doc, "This page manages the clause templates used by the local rule engine.")
    add_heading(doc, "4.1 Tabs", 2)
    add_bullets(doc, [
        ("Product tabs: ", "Stock / Fund / VA / Statement."),
        ("Subtype tabs: ", "shown for products with subtypes (Stock has 6; VA has Buy/Sell; Statement has Monthly / BB Invest / Southbound)."),
        ("Language tabs: ", "English / 繁体中文 / 简体中文."),
    ])
    add_para(doc, "Each product stores clauses per subtype and language, plus a version label.")
    add_heading(doc, "4.2 Manual editing", 2)
    add_bullets(doc, [
        "Click + 添加条款 to add a clause, edit inline, then click 保存全部条款.",
        "Clause numbering is auto-renumbered on save.",
    ])
    add_heading(doc, "4.3 Import from HTML templates", 2)
    add_para(doc, "Three ways to import clauses from HTML templates:")
    add_bullets(doc, [
        ("Single file (单文件导入) — ", "replaces the current product + language clauses."),
        ("Batch (批量导入) — ", "select multiple HTML files; language (_hk_ → zh_hk, _zh_ → zh_cn) and subtype "
                                 "(bbInvest / southBound / default monthly) are detected automatically from the filename."),
        ("Folder sync (一键同步) — ", "scan a whole template folder and import every HTML file, auto-detecting "
                                  "product, subtype and language from the folder path."),
    ])
    add_note(doc, "After importing, click 保存全部条款 to persist to data/clauses.json.")
    add_heading(doc, "4.4 Hints for batch import", 2)
    add_bullets(doc, [
        "Filenames must contain _hk_ / _zh_ to identify the language.",
        "bbInvest / southBound keywords in the filename pick the Statement subtype; otherwise monthly is used.",
        "The version is extracted from a YYMMDD_vN or YYYYMMDD_vN suffix (e.g. ..._240910_v2.html → 2024-09-10 v2).",
    ])

    # 5 Audit engine details -------------------------------------------
    add_heading(doc, "5. Audit Engine Details", 1)
    add_para(doc, "The AI audit (Stock / Fund) combines a deterministic local rule engine with an optional LLM review:")
    add_heading(doc, "5.1 Rule-based checks (local, no API)", 2)
    add_bullets(doc, [
        ("Fee waiver — ", "flags when commission or platform fee is 0, and whether both are waived (zero-fee trade)."),
        ("Compliance — ", "flags fractional-share (<1 share) and micro-amount (<1 currency unit) trades."),
        ("Anomaly — ", "amount consistency (total + commission + platform fee = settlement, 0.02 tolerance) and "
                       "contract-to-settlement date interval (≤0 or >3 days)."),
        ("Data quality — ", "missing order reference or account number."),
    ])
    add_heading(doc, "5.2 AI review (optional)", 2)
    add_para(doc, "When OPENROUTER_API_KEY is set, the engine calls an LLM (model "
                  "nvidia/nemotron-3-nano-30b-a3b:free via OpenRouter) and returns a 3–5 sentence Chinese "
                  "summary focused on fee waiver, compliance and anomalies.")

    # 6 History / 7 chat / 8 export ------------------------------------
    add_heading(doc, "6. History (历史记录)", 1)
    add_para(doc, "Lists all previous audit reports (latest 50 shown). Click a record to view details, or click "
                  "导出 Excel to download a full report.")
    add_heading(doc, "7. AI Assistant (AI 助手)", 1)
    add_para(doc, "The floating chat box in the bottom corner answers questions about file content and audit "
                  "results. It requires OPENROUTER_API_KEY.")
    add_heading(doc, "8. Excel Export", 1)
    add_bullets(doc, [
        ("Each product page: ", "导出 Excel 报告 downloads a clause-audit spreadsheet (clause_audit_report.xlsx)."),
        ("History page: ", "导出 Excel downloads a combined report with three sheets (contract-note summary, "
                            "VA summary, VA detailed checks)."),
    ])

    # 9 FAQ ------------------------------------------------------------
    add_heading(doc, "9. FAQ", 1)
    add_para(doc, "Q: Why is a file tagged 条款问题 even though I imported the template?", bold=True)
    add_para(doc, "Confirm you uploaded the file on the correct product page. A Fund file uploaded on the VA page "
                  "will fail all VA checks.")
    add_para(doc, "Q: What does 跳过 AI 审核 do?", bold=True)
    add_para(doc, "It disables the LLM review, keeping only the fast local rule engine.")
    add_para(doc, "Q: What is the upload limit?", bold=True)
    add_para(doc, "200 PDF files per batch.")
    add_para(doc, "Q: Does the platform work without an API key?", bold=True)
    add_para(doc, "Yes — clause and VA checks run fully offline. Only the AI review and chat require "
                  "OPENROUTER_API_KEY.")

    doc.save(EN_PATH)
    print("wrote", EN_PATH)


# =====================================================================
# CHINESE
# =====================================================================
def build_zh():
    doc = Document()
    set_base_style(doc)

    add_heading(doc, "成交单审核平台 — 用户指南", 0)
    add_para(doc, "一个用于香港投资成交单的 AI 辅助审核平台。平台将成交单与官方模板进行对照，"
                  "找出缺失或不一致的条款。", color=GREY)

    # 1 概览 ------------------------------------------------------------
    add_heading(doc, "1. 平台概览", 1)
    add_para(doc, "平台支持四种产品，每种产品有独立页面，并拥有各自的自动识别与条款核查逻辑。")
    add_table(doc,
              ["产品", "页面", "说明"],
              [
                  ["股票", "Stock Contract Note", "股票交易；自动识别 HK/US 市场及 6 种子类型。"],
                  ["基金", "Fund Contract Note", "基金交易；自动识别语言及 6 种交易子类型。"],
                  ["VA 虚拟资产", "VA Virtual Asset", "虚拟资产成交单；自动识别 Buy/Sell 与语言；额外执行 50 项 VA 专项检查。"],
                  ["投资月结单", "Investment Monthly Statement", "月结单；3 种子类型：月结单 / BB Invest / 南向通。"],
              ],
              widths=[1.4, 2.2, 3.0])
    add_para(doc, "")
    add_para(doc, "每次审核执行以下检查：", bold=True)
    add_bullets(doc, [
        ("条款核查 — ", "本地规则引擎，将 PDF 提取文本与已保存的条款模板和必填字段进行比对。"),
        ("AI 审核 — ", "可选的 LLM 复核（仅股票与基金），标记风险发现、费用减免及总结。"),
        ("VA 专项检查 — ", "虚拟资产成交单专用检查器（50 项）。"),
    ])

    # 2 安装 ------------------------------------------------------------
    add_heading(doc, "2. 安装与启动", 1)
    add_heading(doc, "2.1 环境要求", 2)
    add_bullets(doc, [
        "Python 3.10 或更高版本。",
        "依赖见 requirements.txt（FastAPI、uvicorn、pdfplumber、PyPDF2、pydantic、beautifulsoup4、openpyxl、python-docx、httpx）。",
    ])
    add_heading(doc, "2.2 安装与运行", 2)
    add_code(doc, "pip install -r requirements.txt\ncp .env.example .env    # 然后编辑 .env\npython app.py")
    add_para(doc, "浏览器打开 http://localhost:8000。")
    add_heading(doc, "2.3 配置（.env）", 2)
    add_table(doc,
              ["变量", "用途", "是否必填"],
              [
                  ["OPENROUTER_API_KEY", "AI 审核（OpenRouter/Claude）的 API Key。不设置则仅运行本地规则引擎。", "否"],
                  ["PORT", "服务端口（默认 8000）。", "否"],
              ],
              widths=[1.8, 3.7, 1.0])
    add_para(doc, "")
    add_para(doc, "右上角徽标会显示 AI 是否已启用。没有 API Key 时，条款核查与 VA 检查仍可完全离线运行。")

    # 3 审核 ------------------------------------------------------------
    add_heading(doc, "3. 审核产品", 1)
    add_para(doc, "所有产品页（股票 / 基金 / VA / 月结单）流程一致。")
    add_heading(doc, "第 1 步 — 上传文件", 2)
    add_bullets(doc, [
        "将文件夹或 PDF 文件拖入上传区，或点击 选择 PDF 文件 / 选择文件夹。",
        "支持多文件（单次最多 200 份）。",
        "文件列表会显示已选文件及总大小。",
    ])
    add_heading(doc, "第 2 步 —（仅月结单）选择子类型", 2)
    add_para(doc, "在月结单页面选择 子类型：投资月结单 / BB Invest / Southbound 南向通。")
    add_heading(doc, "第 3 步 —（仅股票/基金）可选跳过 AI", 2)
    add_para(doc, "勾选 跳过 AI 审核，仅运行本地规则引擎。大批量跑量时速度更快。")
    add_heading(doc, "第 4 步 — 开始审核", 2)
    add_para(doc, "点击 开始审核。进度条显示实时上传/处理进度，文件并发处理（最多 10 个线程）。")
    add_heading(doc, "第 5 步 — 查看结果", 2)
    add_para(doc, "每份文件会标注识别出的产品、子类型及总体结果：")
    add_bullets(doc, [
        ("条款通过 — ", "所有条款检查通过。"),
        ("条款问题 — ", "列出未通过的条款。"),
    ])
    add_para(doc, "展开结果可查看：", bold=True)
    add_bullets(doc, [
        ("英文条款 / 中文条款 — ", "逐条通过/失败及匹配关键词。"),
        ("个人资料 — ", "固定字段标签检查。"),
        ("AI 审核 — ", "发现项、严重/警告/提示数量及总结（仅股票/基金）。"),
        ("VA 专项检查 — ", "50 项 VA 清单（仅 VA）。"),
    ])

    add_heading(doc, "3.1 文件类型自动识别", 2)
    add_para(doc, "审核引擎按关键字识别文件类型，因此放错页面的文件也会按内容处理（但仍建议在对应产品页上传）。")
    add_bullets(doc, [
        ("股票子类型 — ", "order（美股买卖）、orderHk（港股买卖）、companyAction / companyActionHk（公司行动通知书）、"
                          "business（资金/股票转账、存入）、businessHk（IPO 认购/分配）。"),
        ("基金子类型 — ", "apply（认购）、redeem（赎回）、deposit（存入）、takeout（提取）、"
                          "divid（现金派息）、divid_share（单位派息）。"),
        ("VA — ", "买入/卖出及语言（繁体/简体）。"),
        ("月结单 — ", "南向通按关键字识别；其余默认 monthly。"),
    ])

    add_heading(doc, "3.2 日期校验", 2)
    add_para(doc, "引擎校验「成交/交易日期 + N 个香港工作日 = 发出日期」，使用香港公众假期表（2024–2027，周末 + 公众假期）。")
    add_bullets(doc, [
        ("股票 / VA — ", "允许 T、T+1 或 T+2 工作日。"),
        ("基金 / 月结单 — ", "严格 T+2 工作日。"),
    ])

    # 4 条款模板 --------------------------------------------------------
    add_heading(doc, "4. 条款模板管理", 1)
    add_para(doc, "本页管理本地规则引擎使用的条款模板。")
    add_heading(doc, "4.1 标签页", 2)
    add_bullets(doc, [
        ("产品标签：", "Stock / Fund / VA / Statement。"),
        ("子类型标签：", "有子类型的产品会显示（股票 6 种；VA 有买入/卖出；月结单有月结单/BB Invest/南向通）。"),
        ("语言标签：", "English / 繁体中文 / 简体中文。"),
    ])
    add_para(doc, "每个产品按子类型与语言分别存储条款，并附带版本号。")
    add_heading(doc, "4.2 手动编辑", 2)
    add_bullets(doc, [
        "点击 + 添加条款 新增条款，内联编辑后点击 保存全部条款。",
        "保存时条款编号会自动重新排序。",
    ])
    add_heading(doc, "4.3 从 HTML 模板导入", 2)
    add_para(doc, "三种方式：")
    add_bullets(doc, [
        ("单文件导入 — ", "替换当前产品 + 语言的条款。"),
        ("批量导入 — ", "选择多个 HTML 文件；从文件名自动识别语言（_hk_ → zh_hk、_zh_ → zh_cn）和子类型"
                       "（bbInvest / southBound / 默认 monthly），并写入对应位置。"),
        ("一键同步 — ", "扫描整个模板文件夹并导入所有 HTML 文件，从路径自动识别产品、子类型与语言。"),
    ])
    add_note(doc, "导入后请点击 保存全部条款，写入 data/clauses.json 生效。")
    add_heading(doc, "4.4 批量导入命名提示", 2)
    add_bullets(doc, [
        "文件名需包含 _hk_ / _zh_ 以识别语言。",
        "文件名中的 bbInvest / southBound 关键字用于识别月结单子类型，否则默认为 monthly。",
        "版本号从 YYMMDD_vN 或 YYYYMMDD_vN 后缀提取（如 ..._240910_v2.html → 2024-09-10 v2）。",
    ])

    # 5 审核引擎 --------------------------------------------------------
    add_heading(doc, "5. 审核引擎细节", 1)
    add_para(doc, "AI 审核（股票/基金）结合确定性的本地规则引擎与可选的 LLM 复核：")
    add_heading(doc, "5.1 规则引擎检查（本地，无需 API）", 2)
    add_bullets(doc, [
        ("费用豁免 — ", "标记佣金或平台费为 0 的情况，以及两项费用是否均豁免（零费用交易）。"),
        ("合规性 — ", "标记小数股（<1 股）与极小金额（<1 单位货币）交易。"),
        ("异常检测 — ", "金额一致性（交易金额 + 佣金 + 平台费 = 交收金额，容差 0.02）及成交-交收日期间隔（≤0 或 >3 天）。"),
        ("数据质量 — ", "缺少交易编号或账户号码。"),
    ])
    add_heading(doc, "5.2 AI 复核（可选）", 2)
    add_para(doc, "设置 OPENROUTER_API_KEY 后，引擎调用 LLM（模型 nvidia/nemotron-3-nano-30b-a3b:free，"
                  "经 OpenRouter）返回 3–5 句中文摘要，重点说明费用豁免、合规性与异常情况。")

    # 6/7/8 ------------------------------------------------------------
    add_heading(doc, "6. 历史记录", 1)
    add_para(doc, "列出所有历史审核报告（显示最近 50 条）。点击记录查看详情，或点击 导出 Excel 下载完整报告。")
    add_heading(doc, "7. AI 助手", 1)
    add_para(doc, "右下角悬浮聊天框可回答关于文件内容与审核结果的问题。需要 OPENROUTER_API_KEY。")
    add_heading(doc, "8. Excel 导出", 1)
    add_bullets(doc, [
        ("各产品页：", "导出 Excel 报告 下载条款审核表（clause_audit_report.xlsx）。"),
        ("历史记录页：", "导出 Excel 下载含三个工作表（成交单汇总、VA 汇总、VA 详细检查项）的综合报告。"),
    ])

    # 9 FAQ ------------------------------------------------------------
    add_heading(doc, "9. 常见问题", 1)
    add_para(doc, "问：为什么明明导入了模板，文件仍显示「条款问题」？", bold=True)
    add_para(doc, "请确认在正确的产品页上传了文件。例如将基金文件上传到 VA 页面会全部失败——应上传到对应产品页。")
    add_para(doc, "问：「跳过 AI 审核」有什么作用？", bold=True)
    add_para(doc, "关闭 LLM 复核，仅保留快速的本地规则引擎。")
    add_para(doc, "问：上传上限是多少？", bold=True)
    add_para(doc, "单次最多 200 份 PDF。")
    add_para(doc, "问：没有 API Key 能否使用？", bold=True)
    add_para(doc, "可以——条款核查与 VA 检查完全离线运行。只有 AI 复核与聊天需要 OPENROUTER_API_KEY。")

    doc.save(ZH_PATH)
    print("wrote", ZH_PATH)


if __name__ == "__main__":
    build_en()
    build_zh()